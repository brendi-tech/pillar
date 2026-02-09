"""
Cloud storage provider for S3/GCS document sync.

Handles the 'cloud_storage' source type by:
1. Listing objects in user's bucket (with optional prefix filter)
2. Downloading and extracting text from supported file types
3. Returning documents as RawItems

Key principle: Files stay in user's bucket - we only read and extract text.
"""
import io
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING, Optional

from apps.knowledge.services.providers.base import (
    BaseProvider,
    RawItem,
    ValidationResult,
)

if TYPE_CHECKING:
    from apps.knowledge.models import KnowledgeSource

logger = logging.getLogger(__name__)


# Supported file types for text extraction
# Matches capabilities of common/services/document_processor.py
SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".md", ".txt",
    ".json", ".pptx", ".ppt", ".xlsx", ".xls",
    ".csv", ".html", ".htm", ".epub",
}

# Maximum file size to process (100MB)
MAX_FILE_SIZE = 100 * 1024 * 1024


@dataclass
class BucketObject:
    """Represents an object in a cloud storage bucket."""

    key: str
    """Object key (path)"""

    etag: str
    """Object etag (for change detection)"""

    size: int
    """Object size in bytes"""

    last_modified: str
    """Last modified timestamp (ISO format)"""


class CloudStorageProvider(BaseProvider):
    """
    Provider for syncing documents from user's cloud storage.

    Supports:
    - Amazon S3
    - Google Cloud Storage

    connection_config schema:
    {
        "provider": "s3",  # or "gcs"
        "bucket": "acme-docs",
        "prefix": "docs/",  # Optional path filter
        "region": "us-east-1",  # S3 only
        "access_key": "encrypted:...",
        "secret_key": "encrypted:...",
        # OR for GCS:
        "credentials_json": "encrypted:...",
    }

    Key behavior:
    - Files stay in user's bucket (no copying)
    - Extract text on sync, store in KnowledgeItem
    - Use etag for change detection on re-sync
    """

    source_type = "cloud_storage"

    async def validate_config(self, source: "KnowledgeSource") -> ValidationResult:
        """
        Verify bucket access.

        Args:
            source: KnowledgeSource with connection_config

        Returns:
            ValidationResult indicating if bucket is accessible
        """
        config = source.connection_config

        # Validate required fields
        provider = config.get("provider")
        bucket = config.get("bucket")

        if not provider:
            return ValidationResult(valid=False, error="Provider not specified")
        if provider not in ("s3", "gcs"):
            return ValidationResult(
                valid=False, error=f"Unsupported provider: {provider}"
            )
        if not bucket:
            return ValidationResult(valid=False, error="Bucket name not specified")

        # Validate bucket name format before attempting connection
        bucket_error = self._validate_bucket_name(bucket, provider)
        if bucket_error:
            return ValidationResult(valid=False, error=bucket_error)

        try:
            client = self._get_client(config)
            # Try to list (limited) to verify access
            await client.list_objects(prefix=config.get("prefix", ""), max_keys=1)
            return ValidationResult(valid=True)
        except Exception as e:
            logger.error(f"Bucket access validation failed: {e}")
            error_msg = self._get_friendly_error_message(str(e), provider)
            return ValidationResult(valid=False, error=error_msg)

    def _validate_bucket_name(self, bucket: str, provider: str) -> str | None:
        """
        Validate bucket name format and return user-friendly error if invalid.

        Args:
            bucket: Bucket name to validate
            provider: 's3' or 'gcs'

        Returns:
            Error message string if invalid, None if valid
        """
        import re

        # Check for leading/trailing whitespace
        if bucket != bucket.strip():
            return "Bucket name cannot have leading or trailing spaces. Please remove any extra spaces."

        # Check for spaces in the middle
        if " " in bucket:
            return "Bucket name cannot contain spaces. Use dashes (-) or underscores (_) instead."

        # Basic character validation (both S3 and GCS have similar rules)
        if provider == "s3":
            # S3 bucket naming rules
            if not re.match(r"^[a-z0-9][a-z0-9.\-]{1,61}[a-z0-9]$", bucket):
                if any(c.isupper() for c in bucket):
                    return "S3 bucket names must be lowercase. Please check for uppercase letters."
                if len(bucket) < 3:
                    return "S3 bucket names must be at least 3 characters long."
                if len(bucket) > 63:
                    return "S3 bucket names cannot exceed 63 characters."
                return "Invalid bucket name format. S3 bucket names must be lowercase and can only contain letters, numbers, hyphens, and periods."
        elif provider == "gcs":
            # GCS bucket naming rules (similar but allows underscores)
            if not re.match(r"^[a-z0-9][a-z0-9._-]{1,61}[a-z0-9]$", bucket):
                if any(c.isupper() for c in bucket):
                    return "GCS bucket names must be lowercase. Please check for uppercase letters."
                if len(bucket) < 3:
                    return "GCS bucket names must be at least 3 characters long."
                if len(bucket) > 63:
                    return "GCS bucket names cannot exceed 63 characters."
                return "Invalid bucket name format. GCS bucket names must be lowercase and can only contain letters, numbers, hyphens, underscores, and periods."

        return None

    def _get_friendly_error_message(self, error: str, provider: str) -> str:
        """
        Convert technical error messages to user-friendly messages.

        Args:
            error: Original error message
            provider: 's3' or 'gcs'

        Returns:
            User-friendly error message
        """
        error_lower = error.lower()

        # Common error patterns
        if "invalid bucket name" in error_lower or "bucket name must match" in error_lower:
            return "Invalid bucket name. Bucket names must be lowercase, 3-63 characters, and can only contain letters, numbers, hyphens, and periods."

        if "access denied" in error_lower or "accessdenied" in error_lower:
            if provider == "s3":
                return "Access denied. Please check that your IAM user has s3:ListBucket and s3:GetObject permissions for this bucket."
            else:
                return "Access denied. Please check that your service account has Storage Object Viewer role for this bucket."

        if "nosuchbucket" in error_lower or "bucket does not exist" in error_lower or "not found" in error_lower:
            return f"Bucket not found. Please verify the bucket name is correct and exists in your {provider.upper()} account."

        if "invalid security token" in error_lower or "security token" in error_lower:
            return "Invalid credentials. The access key or secret key appears to be incorrect."

        if "signaturemismatch" in error_lower or "signature" in error_lower:
            return "Authentication failed. Please verify your secret access key is correct."

        if "invalidaccesskeyid" in error_lower:
            return "Invalid Access Key ID. Please verify your access key is correct and active."

        if "expired" in error_lower:
            return "Credentials have expired. Please generate new access keys."

        if "could not be parsed" in error_lower or "invalid json" in error_lower:
            return "Invalid service account JSON. Please ensure you've pasted the complete JSON key file."

        if "region" in error_lower and ("incorrect" in error_lower or "mismatch" in error_lower):
            return "Region mismatch. Please verify the AWS region matches where your bucket is located."

        if "timeout" in error_lower or "timed out" in error_lower:
            return "Connection timed out. Please try again or check your network connection."

        # If we can't match a specific pattern, return a cleaned-up version
        # but don't expose raw regex patterns
        if "regex" in error_lower or "must match" in error_lower:
            return "Invalid bucket name format. Please check the bucket name for typos, spaces, or invalid characters."

        # Fallback: return original but prefix with "Connection failed"
        return f"Connection failed: {error}"

    async def fetch_items(self, source: "KnowledgeSource") -> list[RawItem]:
        """
        List and extract text from supported documents in the bucket.

        Args:
            source: KnowledgeSource with connection_config

        Returns:
            List of RawItem, one per supported document
        """
        config = source.connection_config
        client = self._get_client(config)

        logger.info(
            f"Listing objects in {config['provider']}://{config['bucket']}"
            f"/{config.get('prefix', '')}"
        )

        objects = await client.list_objects(prefix=config.get("prefix", ""))

        items = []
        for obj in objects:
            # Skip unsupported file types
            if not self._is_supported(obj.key):
                continue

            # Skip files that are too large
            if obj.size > MAX_FILE_SIZE:
                logger.warning(
                    f"Skipping large file: {obj.key} ({obj.size / 1024 / 1024:.1f}MB)"
                )
                continue

            try:
                item = await self._process_object(client, obj, config)
                if item:
                    items.append(item)
            except Exception as e:
                logger.error(f"Error processing {obj.key}: {e}")
                continue

        logger.info(
            f"Fetched {len(items)} documents from {config['provider']}://{config['bucket']}"
        )
        return items

    async def _process_object(
        self, client: "BucketClient", obj: BucketObject, config: dict
    ) -> Optional[RawItem]:
        """
        Download and extract text from a bucket object.

        Args:
            client: Bucket client instance
            obj: BucketObject to process
            config: Source connection config

        Returns:
            RawItem or None if extraction failed
        """
        logger.debug(f"Processing: {obj.key}")

        # Download content
        content_bytes = await client.download_object(obj.key)

        # Extract text
        text = await self._extract_text(content_bytes, obj.key)
        if not text or len(text.strip()) < 10:
            logger.debug(f"Skipping empty document: {obj.key}")
            return None

        # Generate title from filename
        title = Path(obj.key).stem

        return RawItem(
            external_id=obj.etag,  # Use etag for change detection
            title=title,
            raw_content=text,
            url=f"{config['provider']}://{config['bucket']}/{obj.key}",
            metadata={
                "bucket_provider": config["provider"],
                "bucket_name": config["bucket"],
                "object_key": obj.key,
                "object_etag": obj.etag,
                "file_size_bytes": obj.size,
                "last_modified": obj.last_modified,
            },
        )

    def _get_client(self, config: dict) -> "BucketClient":
        """
        Get appropriate bucket client based on provider.

        Args:
            config: Connection config with credentials

        Returns:
            BucketClient instance

        Raises:
            ValueError: If provider is not supported
        """
        provider = config["provider"]

        # Decrypt credentials
        from common.services.credential_encryption import decrypt_value

        if provider == "s3":
            return S3Client(
                bucket=config["bucket"],
                access_key=decrypt_value(config.get("access_key", "")),
                secret_key=decrypt_value(config.get("secret_key", "")),
                region=config.get("region", "us-east-1"),
            )
        elif provider == "gcs":
            return GCSClient(
                bucket=config["bucket"],
                credentials_json=decrypt_value(config.get("credentials_json", "")),
            )
        else:
            raise ValueError(f"Unsupported provider: {provider}")

    def _is_supported(self, key: str) -> bool:
        """Check if file type is supported for text extraction."""
        return Path(key).suffix.lower() in SUPPORTED_EXTENSIONS

    async def _extract_text(self, content: bytes, filename: str) -> str:
        """
        Extract text from document bytes.

        Args:
            content: File content as bytes
            filename: Original filename (for extension detection)

        Returns:
            Extracted text as string
        """
        ext = Path(filename).suffix.lower()

        if ext == ".pdf":
            return await self._extract_pdf(content)
        elif ext in (".docx", ".doc"):
            return await self._extract_docx(content)
        elif ext in (".html", ".htm"):
            return await self._extract_html(content)
        elif ext in (".md", ".txt", ".csv", ".json"):
            return content.decode("utf-8", errors="replace")
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    async def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF bytes."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ""

    async def _extract_docx(self, content: bytes) -> str:
        """Extract text from Word document bytes."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""

    async def _extract_html(self, content: bytes) -> str:
        """Extract text from HTML bytes."""
        try:
            from bs4 import BeautifulSoup

            html_str = content.decode("utf-8", errors="replace")
            soup = BeautifulSoup(html_str, "html.parser")

            # Remove script and style elements
            for element in soup(["script", "style", "nav", "footer", "header"]):
                element.decompose()

            # Get text content
            text = soup.get_text(separator="\n", strip=True)

            # Clean up excessive whitespace
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            return "\n\n".join(lines)
        except Exception as e:
            logger.error(f"HTML extraction failed: {e}")
            return ""


class BucketClient:
    """Base class for bucket clients."""

    async def list_objects(
        self, prefix: str = "", max_keys: int = 1000
    ) -> list[BucketObject]:
        """List objects in the bucket."""
        raise NotImplementedError

    async def download_object(self, key: str) -> bytes:
        """Download object content."""
        raise NotImplementedError


class S3Client(BucketClient):
    """Client for Amazon S3."""

    def __init__(self, bucket: str, access_key: str, secret_key: str, region: str):
        self.bucket = bucket
        self.access_key = access_key
        self.secret_key = secret_key
        self.region = region
        self._client = None

    def _get_client(self):
        """Lazy-load boto3 client."""
        if self._client is None:
            import boto3

            self._client = boto3.client(
                "s3",
                aws_access_key_id=self.access_key,
                aws_secret_access_key=self.secret_key,
                region_name=self.region,
            )
        return self._client

    async def list_objects(
        self, prefix: str = "", max_keys: int = 1000
    ) -> list[BucketObject]:
        """List objects in bucket with optional prefix filter."""
        import asyncio

        def _list():
            client = self._get_client()
            objects = []
            paginator = client.get_paginator("list_objects_v2")

            for page in paginator.paginate(
                Bucket=self.bucket, Prefix=prefix, PaginationConfig={"MaxItems": max_keys}
            ):
                for obj in page.get("Contents", []):
                    objects.append(
                        BucketObject(
                            key=obj["Key"],
                            etag=obj["ETag"].strip('"'),
                            size=obj["Size"],
                            last_modified=obj["LastModified"].isoformat(),
                        )
                    )
            return objects

        return await asyncio.get_event_loop().run_in_executor(None, _list)

    async def download_object(self, key: str) -> bytes:
        """Download object content."""
        import asyncio

        def _download():
            client = self._get_client()
            response = client.get_object(Bucket=self.bucket, Key=key)
            return response["Body"].read()

        return await asyncio.get_event_loop().run_in_executor(None, _download)


class GCSClient(BucketClient):
    """Client for Google Cloud Storage."""

    def __init__(self, bucket: str, credentials_json: str):
        self.bucket_name = bucket
        self.credentials_json = credentials_json
        self._bucket = None

    def _get_bucket(self):
        """Lazy-load GCS bucket."""
        if self._bucket is None:
            import json
            from google.cloud import storage
            from google.oauth2 import service_account

            credentials = service_account.Credentials.from_service_account_info(
                json.loads(self.credentials_json)
            )
            client = storage.Client(credentials=credentials)
            self._bucket = client.bucket(self.bucket_name)
        return self._bucket

    async def list_objects(
        self, prefix: str = "", max_keys: int = 1000
    ) -> list[BucketObject]:
        """List objects in bucket."""
        import asyncio

        def _list():
            bucket = self._get_bucket()
            objects = []
            blobs = bucket.list_blobs(prefix=prefix, max_results=max_keys)

            for blob in blobs:
                objects.append(
                    BucketObject(
                        key=blob.name,
                        etag=blob.etag or "",
                        size=blob.size or 0,
                        last_modified=blob.updated.isoformat() if blob.updated else "",
                    )
                )
            return objects

        return await asyncio.get_event_loop().run_in_executor(None, _list)

    async def download_object(self, key: str) -> bytes:
        """Download object content."""
        import asyncio

        def _download():
            bucket = self._get_bucket()
            blob = bucket.blob(key)
            return blob.download_as_bytes()

        return await asyncio.get_event_loop().run_in_executor(None, _download)


def get_cloud_storage_provider() -> CloudStorageProvider:
    """Get a CloudStorageProvider instance."""
    return CloudStorageProvider()
