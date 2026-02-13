"""
Document upload provider for direct file uploads.

Handles the 'document_upload' source type by:
1. Storing uploaded files in Pillar's GCS bucket
2. Extracting text from supported file types
3. Creating KnowledgeItems for immediate processing

Key principle: Files are stored in Pillar's GCS bucket (not user's).
This is different from cloud_storage which reads from user's bucket.
"""
import asyncio
import hashlib
import io
import logging
from pathlib import Path
from typing import TYPE_CHECKING, Optional

from django.conf import settings
from django.core.files.uploadedfile import UploadedFile

from apps.knowledge.services.providers.base import (
    BaseProvider,
    RawItem,
    ValidationResult,
)

if TYPE_CHECKING:
    from apps.knowledge.models import KnowledgeSource, KnowledgeItem

logger = logging.getLogger(__name__)


# Supported file types for upload
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".md", ".txt"}

# Maximum file size for upload (100MB)
MAX_FILE_SIZE = 100 * 1024 * 1024


class DocumentUploadProvider(BaseProvider):
    """
    Provider for directly uploaded documents.

    Files are stored in Pillar's GCS bucket, not the user's.
    This is the simplest way for users to add documents to their
    knowledge base without setting up cloud storage.

    Storage pattern:
    - documents/{organization_id}/{source_id}/{filename}

    Unlike other providers, DocumentUploadProvider has a special
    handle_upload() method for immediate upload processing, rather
    than periodic sync via fetch_items().
    """

    source_type = "document_upload"

    async def validate_config(self, source: "KnowledgeSource") -> ValidationResult:
        """
        Validate document upload source configuration.

        For document uploads, there's no external configuration needed -
        we just verify the source is properly configured.

        Args:
            source: KnowledgeSource to validate

        Returns:
            ValidationResult (always valid for document_upload type)
        """
        if source.source_type != "document_upload":
            return ValidationResult(
                valid=False, error="Source type must be 'document_upload'"
            )
        return ValidationResult(valid=True)

    async def fetch_items(self, source: "KnowledgeSource") -> list[RawItem]:
        """
        Fetch items from Pillar's GCS bucket for this source.

        Unlike other providers where fetch_items is the main entry point,
        document uploads are typically processed immediately via handle_upload().

        This method is mainly useful for:
        - Re-processing existing uploads (if extraction improves)
        - Rebuilding items if they were deleted

        Args:
            source: KnowledgeSource to fetch from

        Returns:
            List of RawItem from stored documents
        """
        from django.core.files.storage import default_storage

        prefix = f"documents/{source.organization_id}/{source.id}/"

        try:
            # List files in the source's directory
            _, files = default_storage.listdir(prefix.rstrip("/"))
        except Exception as e:
            logger.warning(f"Could not list files for source {source.id}: {e}")
            return []

        items = []
        for filename in files:
            file_path = f"{prefix}{filename}"

            try:
                item = await self._process_stored_file(file_path, filename)
                if item:
                    items.append(item)
            except Exception as e:
                logger.error(f"Error processing stored file {file_path}: {e}")
                continue

        logger.info(f"Fetched {len(items)} documents from storage for source {source.id}")
        return items

    async def _process_stored_file(
        self, file_path: str, filename: str
    ) -> Optional[RawItem]:
        """
        Process a file stored in Django's default storage.

        Args:
            file_path: Full path in storage
            filename: Original filename

        Returns:
            RawItem or None if processing failed
        """
        from django.core.files.storage import default_storage

        # Read file content
        with default_storage.open(file_path, "rb") as f:
            content_bytes = f.read()

        # Extract text
        text = await self._extract_text(content_bytes, filename)
        if not text or len(text.strip()) < 10:
            logger.debug(f"Skipping empty document: {file_path}")
            return None

        # Generate external_id from file path
        external_id = hashlib.md5(file_path.encode()).hexdigest()

        return RawItem(
            external_id=external_id,
            title=Path(filename).stem,
            raw_content=text,
            url=None,  # No external URL for uploads
            metadata={
                "file_path": file_path,
                "original_filename": filename,
                "file_size_bytes": len(content_bytes),
                "file_type": Path(filename).suffix.lstrip("."),
            },
        )

    @classmethod
    def handle_upload(
        cls,
        source: "KnowledgeSource",
        file: UploadedFile,
    ) -> "KnowledgeItem":
        """
        Handle a new file upload (synchronous).

        Saves the file to storage and creates a KnowledgeItem with PENDING status.
        Text extraction and processing happen in background via ProcessingService.

        Args:
            source: Parent KnowledgeSource (must be document_upload type)
            file: Django UploadedFile object

        Returns:
            Created KnowledgeItem (with status=PENDING for background processing)

        Raises:
            ValueError: If file type is not supported or file is too large
        """
        from django.core.files.storage import default_storage
        from django.utils import timezone

        from apps.knowledge.models import KnowledgeItem

        # Validate file
        ext = Path(file.name).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        if file.size > MAX_FILE_SIZE:
            raise ValueError(
                f"File too large: {file.size / 1024 / 1024:.1f}MB (max 100MB)"
            )

        # Generate storage path - use original filename for external_id consistency
        file_path = f"documents/{source.organization_id}/{source.id}/{file.name}"

        # Generate external_id from the intended path (not saved path which may have suffix)
        # This ensures re-uploading the same file updates the existing item
        external_id = hashlib.md5(file_path.encode()).hexdigest()

        # Save file to storage (GCS bucket) - may add suffix if file exists
        saved_path = default_storage.save(file_path, file)
        logger.info(f"Saved upload to: {saved_path}")

        # Create or update KnowledgeItem
        # If re-uploading the same file, update the existing item and re-process
        item, created = KnowledgeItem.objects.update_or_create(
            source=source,
            external_id=external_id,
            defaults={
                "organization": source.organization,
                "product": source.product,  # Denormalized from source
                "item_type": KnowledgeItem.ItemType.PAGE,
                "title": Path(file.name).stem,
                "url": "",  # No external URL
                "raw_content": "",  # Extracted in background by ProcessingService
                "status": KnowledgeItem.Status.PENDING,
                "is_active": True,
                "processing_error": None,  # Clear any previous error
                "metadata": {
                    "file_path": saved_path,
                    "original_filename": file.name,
                    "file_size_bytes": file.size,
                    "file_type": ext.lstrip("."),
                    "uploaded_at": timezone.now().isoformat(),
                },
            },
        )

        action = "Created" if created else "Updated"
        logger.info(f"{action} KnowledgeItem {item.id} for uploaded file: {file.name}")

        # Queue for processing
        from common.task_router import TaskRouter

        TaskRouter.execute("knowledge-process-item", item_id=str(item.id))

        return item

    @classmethod
    async def delete_upload(cls, item: "KnowledgeItem") -> bool:
        """
        Delete an uploaded document and its stored file.

        Args:
            item: KnowledgeItem to delete

        Returns:
            True if deletion was successful
        """
        from django.core.files.storage import default_storage

        file_path = item.metadata.get("file_path")
        if file_path:
            try:
                default_storage.delete(file_path)
                logger.info(f"Deleted file: {file_path}")
            except Exception as e:
                logger.error(f"Failed to delete file {file_path}: {e}")

        # Delete the item (this also cascades to chunks)
        await item.adelete()
        return True

    async def _extract_text(self, content: bytes, filename: str) -> str:
        """Extract text from document bytes (instance method)."""
        return await self._extract_text_static(content, filename)

    @staticmethod
    async def _extract_text_static(content: bytes, filename: str) -> str:
        """
        Extract text from document bytes.

        Args:
            content: File content as bytes
            filename: Original filename (for extension detection)

        Returns:
            Extracted text as string
        """
        ext = Path(filename).suffix.lower()

        try:
            if ext == ".pdf":
                return await DocumentUploadProvider._extract_pdf(content)
            elif ext in (".docx", ".doc"):
                return await DocumentUploadProvider._extract_docx(content)
            elif ext in (".md", ".txt"):
                return content.decode("utf-8", errors="replace")
            else:
                logger.warning(f"Unsupported file type: {ext}")
                return ""
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {e}")
            return ""

    @staticmethod
    async def _extract_pdf(content: bytes) -> str:
        """Extract text from PDF bytes (offloaded to thread pool)."""
        def _sync_extract():
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)

        try:
            return await asyncio.to_thread(_sync_extract)
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ""

    @staticmethod
    async def _extract_docx(content: bytes) -> str:
        """Extract text from Word document bytes (offloaded to thread pool)."""
        def _sync_extract():
            from docx import Document

            doc = Document(io.BytesIO(content))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)

        try:
            return await asyncio.to_thread(_sync_extract)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""


def get_document_upload_provider() -> DocumentUploadProvider:
    """Get a DocumentUploadProvider instance."""
    return DocumentUploadProvider()


def validate_upload_file(file: UploadedFile) -> tuple[bool, Optional[str]]:
    """
    Validate an uploaded file before processing.

    Args:
        file: Django UploadedFile to validate

    Returns:
        Tuple of (is_valid, error_message)
    """
    ext = Path(file.name).suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        return False, f"Unsupported file type: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"

    if file.size > MAX_FILE_SIZE:
        return False, f"File too large: {file.size / 1024 / 1024:.1f}MB (max 100MB)"

    return True, None
